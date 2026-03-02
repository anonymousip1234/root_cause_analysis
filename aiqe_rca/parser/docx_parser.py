"""DOCX document parser using python-docx."""

import io
import uuid

import docx

from aiqe_rca.models.evidence import EvidenceElement, SourceType


def parse_docx(filename: str, content: bytes) -> list[EvidenceElement]:
    """Parse a DOCX file into evidence elements.

    Extracts paragraphs and table content.
    """
    elements: list[EvidenceElement] = []
    counter = 0
    doc = docx.Document(io.BytesIO(content))

    # Extract paragraphs
    current_section = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Track section headings for context
        if para.style and para.style.name and "Heading" in para.style.name:
            current_section = text
            continue

        if len(text) < 10:
            continue

        counter += 1
        ref = f"section: {current_section}" if current_section else None
        elements.append(
            EvidenceElement(
                id=f"E-{uuid.uuid5(uuid.NAMESPACE_DNS, f'{filename}-para-{counter}')}",
                source=filename,
                source_type=SourceType.DOCX,
                text_content=text,
                page_ref=ref,
            )
        )

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        headers = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if row_idx == 0:
                headers = cells
                continue
            row_text_parts = []
            for col_idx, cell_text in enumerate(cells):
                if cell_text:
                    header = headers[col_idx] if col_idx < len(headers) and headers[col_idx] else f"Col{col_idx+1}"
                    row_text_parts.append(f"{header}: {cell_text}")
            row_text = "; ".join(row_text_parts)
            if len(row_text) < 10:
                continue
            counter += 1
            elements.append(
                EvidenceElement(
                    id=f"E-{uuid.uuid5(uuid.NAMESPACE_DNS, f'{filename}-t{table_idx}-r{row_idx}')}",
                    source=filename,
                    source_type=SourceType.DOCX,
                    text_content=row_text,
                    page_ref=f"table {table_idx+1} row {row_idx}",
                )
            )

    return elements
